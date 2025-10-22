from docx import Document

class DocumentGenerator:
    def __init__(self):
        pass

    def export_docx(self, df_or_list, filename, include_answers=False):
        doc = Document()
        title = "Interview Questions" if not include_answers else "Answer Key"
        doc.add_heading(title, level=1)

        if isinstance(df_or_list, pd.DataFrame):
            for i, row in enumerate(df_or_list.itertuples(), 1):
                doc.add_paragraph(f"{i}. {row.Question_Text}")
                if include_answers:
                    doc.add_paragraph(f"Answer: {row.Answer}", style="List Bullet")
                    doc.add_paragraph(f"Explanation: {row.Explanation}", style="List Number")
                doc.add_paragraph("")
        else:  # list of questions (if needed)
            for i, q in enumerate(df_or_list, 1):
                doc.add_paragraph(f"{i}. {q}\n")
                
        doc.save(filename)
        print(f"✅ Saved: {filename}")
